"""Parse DOCX structure and apply surgical text edits while preserving Word formatting."""

from __future__ import annotations

import io
import re
from copy import deepcopy
from dataclasses import dataclass
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Length, Pt

BULLET_PREFIX_RE = re.compile(r"^[\u2022\u2023\u25E6\u2043\-\*\u2013\u2014▪◦]\s*")


@dataclass
class DocxBlock:
    index: int
    text: str
    style_name: str
    is_bullet: bool
    is_in_table: bool
    is_blank: bool


@dataclass
class DocxStructure:
    paragraphs: list[str]
    blocks: list[DocxBlock]
    style_profile: list[dict[str, Any]]
    margins: dict[str, Any]


def is_list_paragraph(paragraph) -> bool:
    p_pr = paragraph._p.pPr
    if p_pr is not None and p_pr.numPr is not None:
        return True
    style = (paragraph.style.name or "").lower() if paragraph.style else ""
    return "list" in style or "bullet" in style


def snapshot_paragraph(paragraph) -> dict[str, Any]:
    pf = paragraph.paragraph_format
    style_name = paragraph.style.name if paragraph.style else "Normal"
    return {
        "style_name": style_name,
        "alignment": paragraph.alignment,
        "left_indent": pf.left_indent,
        "right_indent": pf.right_indent,
        "first_line_indent": pf.first_line_indent,
        "space_before": pf.space_before,
        "space_after": pf.space_after,
        "line_spacing": pf.line_spacing,
        "line_spacing_rule": pf.line_spacing_rule,
        "is_bullet": is_list_paragraph(paragraph),
    }


def snapshot_runs(paragraph) -> list[dict[str, Any]]:
    runs = []
    for run in paragraph.runs:
        runs.append(
            {
                "text": run.text,
                "font_name": run.font.name,
                "font_size": run.font.size,
                "bold": run.bold,
                "italic": run.italic,
            }
        )
    return runs


def parse_docx_structure(data: bytes) -> DocxStructure:
    doc = Document(io.BytesIO(data))
    blocks: list[DocxBlock] = []
    profile: list[dict[str, Any]] = []

    idx = 0
    from docx.text.paragraph import Paragraph

    for child in doc.element.body:
        if child.tag == qn("w:p"):
            p = Paragraph(child, doc)
            _append_block(blocks, profile, p, idx, is_in_table=False)
            idx += 1
        elif child.tag == qn("w:tbl"):
            for row in child.findall(qn("w:tr")):
                for cell in row.findall(qn("w:tc")):
                    for p_el in cell.findall(qn("w:p")):
                        p = Paragraph(p_el, doc)
                        _append_block(blocks, profile, p, idx, is_in_table=True)
                        idx += 1

    section = doc.sections[0]
    margins = {
        "top_margin": section.top_margin,
        "bottom_margin": section.bottom_margin,
        "left_margin": section.left_margin,
        "right_margin": section.right_margin,
    }
    texts = [b.text for b in blocks]
    return DocxStructure(paragraphs=texts, blocks=blocks, style_profile=profile, margins=margins)


def _append_block(
    blocks: list[DocxBlock],
    profile: list[dict[str, Any]],
    paragraph,
    index: int,
    *,
    is_in_table: bool,
) -> None:
    text = paragraph.text
    snap = snapshot_paragraph(paragraph)
    snap["runs"] = snapshot_runs(paragraph)
    profile.append(snap)
    blocks.append(
        DocxBlock(
            index=index,
            text=text,
            style_name=snap["style_name"],
            is_bullet=snap["is_bullet"],
            is_in_table=is_in_table,
            is_blank=not text.strip(),
        )
    )


def strip_manual_bullet_prefix(text: str) -> str:
    return BULLET_PREFIX_RE.sub("", text.strip())


def normalize_ai_text(text: str, block: DocxBlock | None) -> str:
    if text is None:
        return ""
    cleaned = str(text).strip()
    if cleaned in {"(blank line)", "(empty)", "[blank]"}:
        return ""
    if block and block.is_bullet:
        cleaned = strip_manual_bullet_prefix(cleaned)
    return cleaned


def numbered_resume_for_prompt(blocks: list[DocxBlock]) -> str:
    lines = []
    for block in blocks:
        display = block.text if block.text else "(blank line)"
        bullet_flag = "yes" if block.is_bullet else "no"
        table_flag = "yes" if block.is_in_table else "no"
        lines.append(
            f"[{block.index}] style={block.style_name} | bullet={bullet_flag} | "
            f"table={table_flag} | {display!r}"
        )
    return "\n".join(lines)


def merge_edits(original: list[str], edits: list[dict[str, Any]], blocks: list[DocxBlock]) -> list[str]:
    result = list(original)
    for item in edits:
        if not isinstance(item, dict):
            continue
        idx = item.get("index")
        if idx is None:
            continue
        idx = int(idx)
        if 0 <= idx < len(result):
            block = blocks[idx] if idx < len(blocks) else None
            result[idx] = normalize_ai_text(item.get("text", ""), block)
    return result


def align_paragraph_count(original: list[str], updated: list[str]) -> list[str]:
    n = len(original)
    if len(updated) == n:
        return updated
    if len(updated) > n:
        return updated[:n]
    padded = list(updated)
    while len(padded) < n:
        padded.append(original[len(padded)] if len(padded) < len(original) else "")
    return padded


def set_paragraph_text_preserve(paragraph, new_text: str) -> None:
    text = new_text or ""

    if not paragraph.runs:
        paragraph.add_run(text)
        return

    # Fix Skills section bold issue
    # Example:
    # Languages & Scripting: Python | JavaScript | TypeScript
    # Only "Languages & Scripting:" should be bold
    if ":" in text:
        label, rest = text.split(":", 1)

        # Clear old runs
        for run in paragraph.runs:
            run.text = ""

        first = paragraph.runs[0]
        first.text = label + ":"
        first.bold = True

        normal_run = paragraph.add_run(rest)
        normal_run.bold = False
        normal_run.font.name = first.font.name
        normal_run.font.size = first.font.size

        return

    # Normal paragraph or bullet
    first = paragraph.runs[0]
    first.text = text

    for run in paragraph.runs[1:]:
        run.text = ""
  

def _iter_paragraphs_in_order(doc: Document):
    from docx.text.paragraph import Paragraph

    for child in doc.element.body:
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            for row in child.findall(qn("w:tr")):
                for cell in row.findall(qn("w:tc")):
                    for p_el in cell.findall(qn("w:p")):
                        yield Paragraph(p_el, doc)


def apply_paragraphs_to_docx(
    template_bytes: bytes,
    paragraphs: list[str],
    blocks: list[DocxBlock] | None = None,
) -> bytes:
    doc = Document(io.BytesIO(template_bytes))
    refs = list(_iter_paragraphs_in_order(doc))
    for i, paragraph in enumerate(refs):
        if i >= len(paragraphs):
            break

        # Do not touch header/contact area
        # This preserves original name, phone/email, and LinkedIn hyperlink exactly
        if i == 0:
            continue
        new_text = paragraphs[i]
        block = blocks[i] if blocks and i < len(blocks) else None
        if block and block.is_bullet:
            new_text = strip_manual_bullet_prefix(new_text)
        set_paragraph_text_preserve(paragraph, new_text)
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def enrich_style_profile(data: bytes) -> tuple[list[str], list[dict[str, Any]], dict[str, Any]]:
    structure = parse_docx_structure(data)
    return structure.paragraphs, structure.style_profile, structure.margins


def _copy_length(value) -> Length | None:
    if value is None:
        return None
    return deepcopy(value)


def apply_snapshot_to_paragraph(paragraph, snap: dict[str, Any]) -> None:
    try:
        paragraph.style = snap["style_name"]
    except (KeyError, ValueError):
        pass
    paragraph.alignment = snap["alignment"]
    pf = paragraph.paragraph_format
    pf.left_indent = _copy_length(snap["left_indent"])
    pf.right_indent = _copy_length(snap["right_indent"])
    pf.first_line_indent = _copy_length(snap["first_line_indent"])
    pf.space_before = _copy_length(snap["space_before"])
    pf.space_after = _copy_length(snap["space_after"])
    pf.line_spacing = snap["line_spacing"]
    pf.line_spacing_rule = snap["line_spacing_rule"]


def _set_run_font(run, snap: dict[str, Any] | None) -> None:
    if not snap:
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        return
    if snap.get("font_name"):
        run.font.name = snap["font_name"]
    if snap.get("font_size") is not None:
        run.font.size = snap["font_size"]
    if snap.get("bold") is not None:
        run.bold = snap["bold"]
    if snap.get("italic") is not None:
        run.italic = snap["italic"]


def rebuild_docx_from_profile(
    paragraphs: list[str],
    style_profile: list[dict[str, Any]],
    *,
    margins: dict[str, Any] | None = None,
) -> bytes:
    doc = Document()
    if margins:
        section = doc.sections[0]
        section.top_margin = margins["top_margin"]
        section.bottom_margin = margins["bottom_margin"]
        section.left_margin = margins["left_margin"]
        section.right_margin = margins["right_margin"]

    for i, text in enumerate(paragraphs):
        snap = style_profile[i] if i < len(style_profile) else style_profile[-1] if style_profile else {}
        runs = snap.get("runs") or []
        run_snap = runs[0] if runs else snap.get("run")
        style_name = snap.get("style_name", "Normal")
        try:
            p = doc.add_paragraph(style=style_name)
        except (KeyError, ValueError):
            p = doc.add_paragraph()
        if snap:
            apply_snapshot_to_paragraph(p, snap)
        if not text:
            continue
        if snap.get("is_bullet"):
            text = strip_manual_bullet_prefix(text)
        run = p.add_run(text)
        _set_run_font(run, run_snap)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# Backward-compatible helpers
def extract_docx_paragraphs(data: bytes) -> list[str]:
    return parse_docx_structure(data).paragraphs


def numbered_resume_for_prompt_legacy(paragraphs: list[str]) -> str:
    blocks = [
        DocxBlock(i, t, "Normal", t.lstrip().startswith(("-", "•", "*")), False, not t.strip())
        for i, t in enumerate(paragraphs)
    ]
    return numbered_resume_for_prompt(blocks)
