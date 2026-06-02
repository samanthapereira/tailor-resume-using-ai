import io
import json
import os
import re
from datetime import datetime
from pathlib import Path
from typing import Any

import streamlit as st
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Inches, Pt
from dotenv import load_dotenv
from openai import OpenAI

from docx_format import (
    BULLET_PREFIX_RE,
    DocxBlock,
    align_paragraph_count,
    apply_paragraphs_to_docx,
    merge_edits,
    numbered_resume_for_prompt,
    parse_docx_structure,
    rebuild_docx_from_profile,
)

BASE_DIR = Path(__file__).resolve().parent
RESUMES_DIR = BASE_DIR / "resumes"
OUTPUTS_DIR = BASE_DIR / "outputs"

load_dotenv(BASE_DIR / ".env", override=True)
print("API KEY LOADED:", os.getenv("OPENAI_API_KEY"))

TAILOR_SYSTEM_PROMPT = """You are an expert senior technical resume optimizer specializing in ATS optimization WITHOUT degrading readability, credibility, or visual flow.

TAILOR the resume to the job description while preserving:
- the original formatting structure and section order
- bullet rhythm and whitespace
- leadership presence and seniority
- quantified achievements and metrics (never remove or weaken numbers)
- executive tone and concise engineering language
- visual readability for recruiters

CRITICAL RULES:
- DO NOT rewrite the entire resume — make targeted, minimal edits only
- DO NOT make the resume sound AI-generated
- DO NOT replace strong quantified bullets with generic keyword-heavy statements
- Preserve metrics, impact, and technical credibility
- Keep bullets concise and skimmable (action + technology + measurable impact)
- Maintain natural human-written engineering language

TAILORING:
- Inject missing ATS keywords naturally into existing bullets where honestly supported
- Add only highly relevant job-description concepts
- Prioritize semantic alignment over keyword stuffing
- Keep strong original accomplishments intact
- Only modify bullets when necessary for ATS alignment
- Avoid repetitive buzzwords and filler (e.g. excessive "cloud-native", "end-to-end", "robust", "innovative")

FACTUAL RULES:
- Never invent employers, titles, dates, degrees, certifications, skills, or metrics
- Do not add skills or tools not demonstrated in the source resume

FORMATTING PRESERVATION (parse DOCX — preserve structure):
- Preserve EXACT formatting: paragraph styles, bullets, indentation, spacing, line breaks, fonts, margins
- DO NOT convert bullets to paragraphs, reflow sections, or redesign layout
- Modify ONLY text content required for ATS — surgical edits inside existing bullets
- Minimal edits; preserve voice, whitespace, and recruiter-safe readability

OUTPUT:
- Return edits only (see format rules); unchanged blocks must not be rewritten"""

FORMAT_RULES_PARAGRAPH = """
CRITICAL — MINIMAL EDITS ONLY ({count} blocks, indices [0]..[{last_index}]):
- Each block shows: style= | bullet=yes/no | table=yes/no | text
- Return JSON key "edits" with ONLY blocks you changed: [{{"index": 5, "text": "new text"}}]
- If nothing needs changing, return {{"edits": []}}
- bullet=yes: do NOT type bullet symbols (•, -, *) — Word renders them; text only
- Do NOT merge, split, reorder, or renumber blocks
- Leave (blank line) blocks out of edits unless absolutely necessary
- Preserve metrics, senior tone, and human-written engineering voice
- Do not return the full "paragraphs" array unless unavoidable
"""

TAILOR_USER_TEMPLATE = """## Base resume
{resume}

## Job description
{job_description}

Return ONLY the tailored resume text."""

TAILOR_PARAGRAPH_USER_TEMPLATE = """## Job description
{job_description}

## Resume paragraphs (preserve structure — edit text only)
{numbered_resume}

{format_rules}

Return JSON only (prefer minimal edits):
{{
  "edits": [{{"index": 12, "text": "revised bullet text only"}}]
}}"""

REVIEW_SYSTEM_PROMPT = """You are a senior engineering hiring manager performing a second-pass review of a tailored resume.

Review for:
- Does it still sound human-written (not AI-generated)?
- Are any bullets too verbose? Shorten them.
- Did leadership presence decrease? Restore it.
- Did readability or scanability worsen? Improve visual rhythm and whitespace.
- Were strong quantified achievements weakened? Restore metrics.
- Are there repeated buzzwords? Remove or vary them.
- Would you believe this resume from an experienced engineer?

If needed: shorten bullets, restore metrics, reduce keyword stuffing, improve executive presence, restore concise technical storytelling.

FACTUAL RULES: Never invent experience, skills, or metrics. Do not change employers, dates, or titles.

FORMATTING: Same paragraph count and order as input. Do not merge/split paragraphs. Preserve blank lines.

OUTPUT: Return valid JSON only."""

REVIEW_TEXT_USER_TEMPLATE = """## Original base resume
{base_resume}

## Job description
{job_description}

## Tailored resume (to review and improve)
{tailored_resume}

Return JSON:
{{
  "resume": "<full improved resume text, same structure as original>",
  "changes_made": ["<brief note on each meaningful fix>"]
}}"""

REVIEW_PARAGRAPH_USER_TEMPLATE = """## Original base resume
{base_resume}

## Job description
{job_description}

## Resume paragraphs to review
{numbered_resume}

{format_rules}

Return JSON (minimal edits preferred):
{{
  "edits": [{{"index": 12, "text": "..."}}],
  "changes_made": ["<brief note on each meaningful fix>"]
}}"""

SKILLS_ENHANCE_SYSTEM_PROMPT = """You update resume Skills sections for ATS using ONLY skills evidenced in the candidate's original resume.

Rules:
- Never invent skills, tools, or certifications
- Add job-description keywords to Skills only when the base resume clearly supports them
- Keep all non-Skills content unchanged unless a tiny fix is required for consistency
- Preserve exact paragraph count and order; only edit Skills-related paragraph text
- Return valid JSON only"""

SKILLS_TEXT_USER_TEMPLATE = """## Original base resume (source of truth)
{base_resume}

## Job description
{job_description}

## Current tailored resume
{tailored_resume}

Add missing JD keywords to the Skills section only when supported by the base resume.

Return JSON:
{{
  "resume": "<full resume text>",
  "added_skills": ["..."],
  "skipped_skills": [{{"skill": "...", "reason": "..."}}]
}}"""

SKILLS_PARAGRAPH_USER_TEMPLATE = """## Original base resume (source of truth)
{base_resume}

## Job description
{job_description}

## Resume paragraphs
{numbered_resume}

{format_rules}

Add missing JD keywords to Skills paragraph(s) only when supported by the base resume.

Return JSON (only edit Skills-related block indices):
{{
  "edits": [{{"index": 20, "text": "..."}}],
  "added_skills": ["..."],
  "skipped_skills": [{{"skill": "...", "reason": "..."}}]
}}"""

SECTION_HEADERS = {
    "EXPERIENCE",
    "PROFESSIONAL EXPERIENCE",
    "WORK EXPERIENCE",
    "EMPLOYMENT",
    "EDUCATION",
    "SKILLS",
    "TECHNICAL SKILLS",
    "CORE COMPETENCIES",
    "SUMMARY",
    "PROFESSIONAL SUMMARY",
    "PROFILE",
    "CERTIFICATIONS",
    "PROJECTS",
    "LEADERSHIP",
    "AWARDS",
    "PUBLICATIONS",
}


def get_client() -> OpenAI:
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise ValueError("Set OPENAI_API_KEY in .env")
    return OpenAI(api_key=api_key)


def call_openai(system: str, user: str, model: str, *, temperature: float, json_mode: bool = False) -> str:
    kwargs: dict[str, Any] = {
        "model": model,
        "messages": [
            {"role": "system", "content": system},
            {"role": "user", "content": user},
        ],
        "temperature": temperature,
    }
    if json_mode:
        kwargs["response_format"] = {"type": "json_object"}
    response = get_client().chat.completions.create(**kwargs)
    return (response.choices[0].message.content or "").strip()


def extract_text_from_upload(uploaded_file) -> str:
    name = (uploaded_file.name or "").lower()
    data = uploaded_file.getvalue()

    if name.endswith(".txt"):
        return data.decode("utf-8", errors="replace")

    if name.endswith(".pdf"):
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    if name.endswith(".docx"):
        from docx import Document

        doc = Document(io.BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs)

    raise ValueError("Unsupported file type. Use .txt, .pdf, or .docx.")


def _format_rules_block(count: int) -> str:
    return FORMAT_RULES_PARAGRAPH.format(count=count, last_index=count - 1)

def paragraphs_to_text(paragraphs: list[str]) -> str:
    return "\n".join(paragraphs)

def clean_duplicate_words(text: str) -> str:
    text = text.replace("LinkedInLinkedIn", "LinkedIn")
    return text


def restore_header(original: list[str], updated: list[str]) -> list[str]:
    result = list(updated)

    # Keep the first 3 lines exactly like the original resume
    # This protects name, phone/email/location, LinkedIn
    for i, text in enumerate(original[:3]):
        if i < len(result):
            result[i] = clean_duplicate_words(text)

    return result

def _parse_paragraphs_response(
    data: dict[str, Any], original: list[str], blocks: list[DocxBlock]
) -> list[str]:
    edits = data.get("edits")
    if isinstance(edits, list):
        if edits:
            return merge_edits(original, edits, blocks)
        return list(original)

    paragraphs = data.get("paragraphs")
    if isinstance(paragraphs, list):
        texts = [str(p) if p is not None else "" for p in paragraphs]
        texts = [t if t not in {"(blank line)", "(empty)", "[blank]"} else "" for t in texts]
        return align_paragraph_count(original, texts)

    raise ValueError("AI response must include 'edits' or 'paragraphs'.")


def tailor_resume(resume: str, job_description: str, model: str) -> str:
    return call_openai(
        TAILOR_SYSTEM_PROMPT,
        TAILOR_USER_TEMPLATE.format(resume=resume.strip(), job_description=job_description.strip()),
        model,
        temperature=0.35,
    )


def tailor_resume_paragraphs(
    paragraphs: list[str], blocks: list[DocxBlock], job_description: str, model: str
) -> list[str]:
    count = len(paragraphs)
    raw = call_openai(
        TAILOR_SYSTEM_PROMPT,
        TAILOR_PARAGRAPH_USER_TEMPLATE.format(
            job_description=job_description.strip(),
            numbered_resume=numbered_resume_for_prompt(blocks),
            format_rules=_format_rules_block(count),
            count=count,
        ),
        model,
        temperature=0.35,
        json_mode=True,
    )
    data = json.loads(raw)
    return _parse_paragraphs_response(data, paragraphs, blocks)


def review_resume(base_resume: str, tailored: str, job_description: str, model: str) -> dict[str, Any]:
    raw = call_openai(
        REVIEW_SYSTEM_PROMPT,
        REVIEW_TEXT_USER_TEMPLATE.format(
            base_resume=base_resume.strip(),
            job_description=job_description.strip(),
            tailored_resume=tailored.strip(),
        ),
        model,
        temperature=0.3,
        json_mode=True,
    )
    data = json.loads(raw)
    resume = (data.get("resume") or "").strip()
    if not resume:
        raise ValueError("Review pass returned an empty resume.")
    return {
        "resume": resume,
        "changes_made": data.get("changes_made") or [],
    }


def review_resume_paragraphs(
    base_resume: str,
    paragraphs: list[str],
    blocks: list[DocxBlock],
    job_description: str,
    model: str,
) -> dict[str, Any]:
    count = len(paragraphs)
    raw = call_openai(
        REVIEW_SYSTEM_PROMPT,
        REVIEW_PARAGRAPH_USER_TEMPLATE.format(
            base_resume=base_resume.strip(),
            job_description=job_description.strip(),
            numbered_resume=numbered_resume_for_prompt(blocks),
            format_rules=_format_rules_block(count),
            count=count,
        ),
        model,
        temperature=0.3,
        json_mode=True,
    )
    data = json.loads(raw)
    updated = _parse_paragraphs_response(data, paragraphs, blocks)
    return {
        "paragraphs": updated,
        "changes_made": data.get("changes_made") or [],
    }


def enhance_skills_for_ats(base_resume: str, tailored: str, job_description: str, model: str) -> dict[str, Any]:
    raw = call_openai(
        SKILLS_ENHANCE_SYSTEM_PROMPT,
        SKILLS_TEXT_USER_TEMPLATE.format(
            base_resume=base_resume.strip(),
            job_description=job_description.strip(),
            tailored_resume=tailored.strip(),
        ),
        model,
        temperature=0.2,
        json_mode=True,
    )
    data = json.loads(raw)
    resume = (data.get("resume") or "").strip()
    if not resume:
        raise ValueError("Skills enhancement returned an empty resume.")
    return {
        "resume": resume,
        "added_skills": data.get("added_skills") or [],
        "skipped_skills": data.get("skipped_skills") or [],
    }


def enhance_skills_paragraphs(
    base_resume: str,
    paragraphs: list[str],
    blocks: list[DocxBlock],
    job_description: str,
    model: str,
) -> dict[str, Any]:
    count = len(paragraphs)
    raw = call_openai(
        SKILLS_ENHANCE_SYSTEM_PROMPT,
        SKILLS_PARAGRAPH_USER_TEMPLATE.format(
            base_resume=base_resume.strip(),
            job_description=job_description.strip(),
            numbered_resume=numbered_resume_for_prompt(blocks),
            format_rules=_format_rules_block(count),
            count=count,
        ),
        model,
        temperature=0.2,
        json_mode=True,
    )
    data = json.loads(raw)
    updated = _parse_paragraphs_response(data, paragraphs, blocks)
    return {
        "paragraphs": updated,
        "added_skills": data.get("added_skills") or [],
        "skipped_skills": data.get("skipped_skills") or [],
    }


def is_section_header(line: str) -> bool:
    stripped = line.strip()
    if not stripped or len(stripped) > 60:
        return False
    upper = stripped.upper().rstrip(":")
    if upper in SECTION_HEADERS:
        return True
    if stripped.isupper() and len(stripped.split()) <= 5 and not stripped.startswith(("-", "•", "*")):
        return True
    if re.match(r"^[A-Z][A-Z\s/&\-]{2,50}:?\s*$", stripped):
        return True
    return False


def is_contact_line(line: str) -> bool:
    lower = line.lower()
    return any(x in lower for x in ("@", "linkedin", "github", "http", "www.")) or bool(
        re.search(r"\d{3}[-.\s]?\d{3}[-.\s]?\d{4}", line)
    )


def is_job_header_line(line: str) -> bool:
    stripped = line.strip()
    if stripped.startswith(("-", "•", "*")):
        return False
    if "|" in stripped and len(stripped) < 120:
        return True
    if re.search(r"\b(19|20)\d{2}\b", stripped) and len(stripped) < 100:
        return True
    return False


def _set_run_font(run, *, size: int, bold: bool = False) -> None:
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)


def _style_paragraph(paragraph, *, before: int = 0, after: int = 2, line_spacing: float = 1.08) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = line_spacing


def text_to_docx_bytes(content: str) -> bytes:
    from docx import Document

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)

    lines = content.splitlines()
    non_empty = [ln for ln in lines if ln.strip()]
    name_set = False

    for idx, line in enumerate(lines):
        stripped = line.strip()

        if not stripped:
            doc.add_paragraph("")
            continue

        if not name_set and idx < 3 and not is_section_header(stripped) and not stripped.startswith(("-", "•", "*")):
            if not is_contact_line(stripped) and len(stripped) < 80:
                p = doc.add_paragraph()
                p.alignment = 0
                run = p.add_run(stripped)
                _set_run_font(run, size=16, bold=True)
                _style_paragraph(p, before=0, after=4, line_spacing=1.0)
                name_set = True
                continue

        if is_contact_line(stripped):
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            _set_run_font(run, size=10, bold=False)
            _style_paragraph(p, before=0, after=8, line_spacing=1.0)
            continue

        if is_section_header(stripped):
            p = doc.add_paragraph()
            header_text = stripped.rstrip(":").upper()
            run = p.add_run(header_text)
            _set_run_font(run, size=11, bold=True)
            _style_paragraph(p, before=10, after=4, line_spacing=1.0)
            continue

        if stripped.startswith(("- ", "• ", "* ", "– ")):
            bullet_text = re.sub(r"^[-•*–]\s+", "", stripped)
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(bullet_text)
            _set_run_font(run, size=10, bold=False)
            _style_paragraph(p, before=0, after=2, line_spacing=1.08)
            continue

        if is_job_header_line(stripped):
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            _set_run_font(run, size=10, bold=True)
            _style_paragraph(p, before=6, after=2, line_spacing=1.05)
            continue

        p = doc.add_paragraph()
        run = p.add_run(stripped)
        _set_run_font(run, size=10, bold=False)
        _style_paragraph(p, before=0, after=3, line_spacing=1.08)

    if not non_empty:
        p = doc.add_paragraph()
        p.add_run(content)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def sanitize_filename_part(value: str, *, lowercase: bool = False) -> str:
    cleaned = re.sub(r"[^\w\s-]", "", value.strip())
    cleaned = re.sub(r"[\s-]+", "_", cleaned).strip("_")
    if lowercase:
        cleaned = cleaned.lower()
    return cleaned or "unknown"


def build_output_filename(company_name: str, first_name: str) -> str:
    company = sanitize_filename_part(company_name, lowercase=True)
    name_parts = [p.capitalize() for p in sanitize_filename_part(first_name).split("_") if p]
    name = "_".join(name_parts) if name_parts else "Unknown"
    date_part = datetime.now().strftime("%m%d")
    return f"{company}_resume_{name}_{date_part}.docx"


def save_uploaded_resume(uploaded_file) -> Path:
    RESUMES_DIR.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    path = RESUMES_DIR / f"{timestamp}_{uploaded_file.name}"
    path.write_bytes(uploaded_file.getvalue())
    return path


def build_docx_output(
    paragraphs: list[str],
    *,
    template_bytes: bytes | None,
    style_profile: list[dict] | None,
    margins: dict | None,
    fallback_text: str,
    blocks: list[DocxBlock] | None = None,
) -> bytes:
    if template_bytes:
        return apply_paragraphs_to_docx(template_bytes, paragraphs, blocks)
    if style_profile:
        return rebuild_docx_from_profile(paragraphs, style_profile, margins=margins)
    return text_to_docx_bytes(fallback_text)


def save_tailored_output(
    paragraphs: list[str],
    company_name: str,
    first_name: str,
    *,
    template_bytes: bytes | None = None,
    style_profile: list[dict] | None = None,
    margins: dict | None = None,
    fallback_text: str = "",
    blocks: list[DocxBlock] | None = None,
) -> tuple[Path, bytes, str]:
    OUTPUTS_DIR.mkdir(parents=True, exist_ok=True)
    filename = build_output_filename(company_name, first_name)
    docx_bytes = build_docx_output(
        paragraphs,
        template_bytes=template_bytes,
        style_profile=style_profile,
        margins=margins,
        fallback_text=fallback_text or paragraphs_to_text(paragraphs),
        blocks=blocks,
    )
    path = OUTPUTS_DIR / filename
    path.write_bytes(docx_bytes)
    return path, docx_bytes, filename


def main() -> None:
    st.set_page_config(page_title="Resume Tailor AI", page_icon="📄", layout="wide")
    st.title("Resume Tailor AI")
    st.caption(
        "ATS optimization with executive readability. Upload a .docx to preserve your original formatting."
    )

    default_model = os.getenv("OPENAI_MODEL", "gpt-4o-mini")

    with st.sidebar:
        st.header("Settings")
        model = st.text_input("OpenAI model", value=default_model)
        st.markdown("API key is loaded from `.env`")
        st.divider()
        st.subheader("Output file naming")
        company_name = st.text_input("Company name", placeholder="e.g. Acme Corp")
        first_name = st.text_input("Your first name", placeholder="e.g. Samantha")
        st.divider()
        enhance_skills = st.checkbox(
            "Boost Skills for ATS",
            value=True,
            help="Adds JD keywords to Skills when your base resume supports them.",
        )
        second_review = st.checkbox(
            "Second review pass",
            value=True,
            help="Senior hiring-manager review for tone, brevity, metrics, and readability.",
        )

    col1, col2 = st.columns(2)

    with col1:
        st.subheader("Base resume")
        upload = st.file_uploader(
            "Upload resume (.docx required for format preservation)",
            type=["txt", "pdf", "docx"],
            help="Your original .docx is used as the template — fonts, bullets, indentation, and spacing stay intact.",
        )
        resume_text = st.text_area(
            "Or paste resume",
            height=320,
            placeholder="Paste your full resume here…",
        )

    with col2:
        st.subheader("Job description")
        job_description = st.text_area(
            "Paste job posting",
            height=400,
            placeholder="Paste the full job description…",
        )

    if st.button("Tailor resume", type="primary", use_container_width=True):
        resume = resume_text.strip()
        template_bytes: bytes | None = None
        style_profile: list[dict] | None = None
        margins: dict | None = None
        source_paragraphs: list[str] | None = None
        docx_blocks: list[DocxBlock] | None = None
        use_format_preservation = False

        if upload is not None:
            try:
                save_uploaded_resume(upload)
                upload_name = (upload.name or "").lower()
                upload_bytes = upload.getvalue()
                if upload_name.endswith(".docx"):
                    template_bytes = upload_bytes
                    structure = parse_docx_structure(upload_bytes)
                    source_paragraphs = structure.paragraphs
                    docx_blocks = structure.blocks
                    style_profile = structure.style_profile
                    margins = structure.margins
                    use_format_preservation = True
                    resume = paragraphs_to_text(source_paragraphs)
                else:
                    extracted = extract_text_from_upload(upload)
                    resume = extracted if not resume else f"{resume}\n\n---\n\n{extracted}"
            except Exception as exc:
                st.error(str(exc))
                st.stop()

        if not resume:
            st.warning("Provide a resume (upload or paste).")
            st.stop()
        if not job_description.strip():
            st.warning("Paste a job description.")
            st.stop()
        if not company_name.strip():
            st.warning("Enter the company name in the sidebar.")
            st.stop()
        if not first_name.strip():
            st.warning("Enter your first name in the sidebar.")
            st.stop()

        if use_format_preservation and source_paragraphs is not None and docx_blocks is not None:
            working_paragraphs = list(source_paragraphs)
            working_blocks = docx_blocks
        else:
            working_paragraphs = [ln for ln in resume.splitlines()]
            working_blocks = [
                DocxBlock(i, t, "Normal", bool(BULLET_PREFIX_RE.match(t.strip())), False, not t.strip())
                for i, t in enumerate(working_paragraphs)
            ]
            if not template_bytes:
                st.warning(
                    "Upload a **.docx** file to preserve fonts, bullets, indentation, and spacing. "
                    "Paste/PDF uses auto-generated layout."
                )

        try:
            if use_format_preservation:
                with st.spinner("Tailoring (surgical edits, preserving layout)…"):
                    working_paragraphs = tailor_resume_paragraphs(
                        working_paragraphs, working_blocks, job_description, model
                    )
                    working_paragraphs = restore_header(source_paragraphs, working_paragraphs)

                if enhance_skills:
                    with st.spinner("Adding supported keywords to Skills…"):
                        skills_result = enhance_skills_paragraphs(
                            resume, working_paragraphs, working_blocks, job_description, model
                        )
                        working_paragraphs = skills_result["paragraphs"]
                        working_paragraphs = restore_header(source_paragraphs, working_paragraphs)
                        st.session_state["added_skills"] = skills_result["added_skills"]
                        st.session_state["skipped_skills"] = skills_result["skipped_skills"]
                else:
                    st.session_state["added_skills"] = []
                    st.session_state["skipped_skills"] = []

                if second_review:
                    with st.spinner("Second review (readability & executive tone)…"):
                        review_result = review_resume_paragraphs(
                            resume, working_paragraphs, working_blocks, job_description, model
                        )
                        working_paragraphs = review_result["paragraphs"]
                        working_paragraphs = restore_header(source_paragraphs, working_paragraphs)
                        st.session_state["review_changes"] = review_result["changes_made"]
                else:
                    st.session_state["review_changes"] = []
            else:
                tailored = ""
                with st.spinner("Tailoring (minimal edits, preserve structure)…"):
                    tailored = tailor_resume(resume, job_description, model)

                if enhance_skills:
                    with st.spinner("Adding supported keywords to Skills…"):
                        skills_result = enhance_skills_for_ats(
                            resume, tailored, job_description, model
                        )
                        tailored = skills_result["resume"]
                        st.session_state["added_skills"] = skills_result["added_skills"]
                        st.session_state["skipped_skills"] = skills_result["skipped_skills"]
                else:
                    st.session_state["added_skills"] = []
                    st.session_state["skipped_skills"] = []

                if second_review:
                    with st.spinner("Second review (readability & executive tone)…"):
                        review_result = review_resume(resume, tailored, job_description, model)
                        tailored = review_result["resume"]
                        st.session_state["review_changes"] = review_result["changes_made"]
                else:
                    st.session_state["review_changes"] = []

                working_paragraphs = [ln for ln in tailored.splitlines()]

        except json.JSONDecodeError as exc:
            st.error(f"Could not parse AI response: {exc}")
            st.stop()
        except Exception as exc:
            st.error(f"OpenAI error: {exc}")
            st.stop()

        tailored_text = paragraphs_to_text(working_paragraphs)
        output_path, docx_bytes, output_filename = save_tailored_output(
            working_paragraphs,
            company_name,
            first_name,
            template_bytes=template_bytes,
            style_profile=style_profile,
            margins=margins,
            fallback_text=tailored_text,
            blocks=working_blocks if use_format_preservation else None,
        )
        st.session_state["tailored"] = tailored_text
        st.session_state["output_path"] = str(output_path)
        st.session_state["docx_bytes"] = docx_bytes
        st.session_state["output_filename"] = output_filename
        st.session_state["format_preserved"] = use_format_preservation

    if tailored := st.session_state.get("tailored"):
        added = st.session_state.get("added_skills") or []
        skipped = st.session_state.get("skipped_skills") or []
        review_changes = st.session_state.get("review_changes") or []

        if added or skipped or review_changes:
            with st.expander("Optimization report", expanded=bool(added or review_changes)):
                if review_changes:
                    st.markdown("**Second review adjustments**")
                    for note in review_changes:
                        st.caption(f"- {note}")
                if added:
                    st.markdown("**Skills added** (supported by base resume)")
                    st.write(", ".join(str(s) for s in added))
                if skipped:
                    st.markdown("**Skills not added**")
                    for item in skipped:
                        if isinstance(item, dict):
                            st.caption(f"- **{item.get('skill', '?')}**: {item.get('reason', '')}")
                        else:
                            st.caption(f"- {item}")

        st.subheader("Tailored resume")
        st.text_area("Result", value=tailored, height=480, label_visibility="collapsed")

        output_path = st.session_state.get("output_path", "")
        if output_path:
            preserved = st.session_state.get("format_preserved", False)
            note = "original .docx layout preserved" if preserved else "generated layout (upload .docx to preserve yours)"
            st.caption(f"Saved to `{output_path}` — {note}")

        docx_bytes = st.session_state.get("docx_bytes") or text_to_docx_bytes(tailored)
        download_name = st.session_state.get("output_filename", "tailored_resume.docx")
        st.download_button(
            "Download as .docx",
            data=docx_bytes,
            file_name=download_name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )


if __name__ == "__main__":
    main()
